"""
ARIA - Document Generation Tool

Purpose: Generate DOCX, XLSX, and PDF documents from structured content.
"""

import logging
import os
from pathlib import Path

from ..base import BaseTool, ToolParameter, ToolResult, ToolStatus, ToolType
from aria.config import settings

logger = logging.getLogger(__name__)


class DocumentGenerationTool(BaseTool):
    """Generate documents in DOCX, XLSX, or PDF format."""

    @property
    def name(self) -> str:
        return "generate_document"

    @property
    def description(self) -> str:
        return (
            "Generate a document file (DOCX, XLSX, or PDF). "
            "Provide structured content and a filename. "
            "The file is saved to the configured output directory."
        )

    @property
    def type(self) -> ToolType:
        return ToolType.BUILTIN

    @property
    def parameters(self) -> list[ToolParameter]:
        return [
            ToolParameter(
                name="format",
                type="string",
                description="Output format: docx, xlsx, or pdf",
                required=True,
                enum=["docx", "xlsx", "pdf"],
            ),
            ToolParameter(
                name="filename",
                type="string",
                description="Output filename (without extension)",
                required=True,
            ),
            ToolParameter(
                name="content",
                type="object",
                description=(
                    "Document content. For docx: {title, paragraphs: [str], headings: [{level, text}]}. "
                    "For xlsx: {sheets: [{name, headers: [str], rows: [[val]]}]}. "
                    "For pdf: {title, paragraphs: [str]}."
                ),
                required=True,
            ),
        ]

    def _ensure_output_dir(self) -> Path:
        output_dir = Path(os.path.expanduser(settings.docgen_output_dir))
        output_dir.mkdir(parents=True, exist_ok=True)
        return output_dir

    async def execute(self, arguments: dict) -> ToolResult:
        fmt = arguments.get("format", "")
        filename = arguments.get("filename", "document")
        content = arguments.get("content", {})

        if fmt == "docx":
            return await self._generate_docx(filename, content)
        elif fmt == "xlsx":
            return await self._generate_xlsx(filename, content)
        elif fmt == "pdf":
            return await self._generate_pdf(filename, content)
        else:
            return ToolResult(
                tool_name=self.name,
                status=ToolStatus.ERROR,
                error=f"Unsupported format: {fmt}. Use docx, xlsx, or pdf.",
            )

    async def _generate_docx(self, filename: str, content: dict) -> ToolResult:
        try:
            from docx import Document
        except ImportError:
            return ToolResult(
                tool_name=self.name,
                status=ToolStatus.ERROR,
                error="python-docx not installed. Install with: pip install python-docx",
            )

        try:
            doc = Document()

            title = content.get("title")
            if title:
                doc.add_heading(title, level=0)

            for item in content.get("headings", []):
                doc.add_heading(item.get("text", ""), level=item.get("level", 1))

            for para in content.get("paragraphs", []):
                doc.add_paragraph(para)

            output_dir = self._ensure_output_dir()
            filepath = output_dir / f"{filename}.docx"
            doc.save(str(filepath))

            return ToolResult(
                tool_name=self.name,
                status=ToolStatus.SUCCESS,
                output=f"Document saved to {filepath}",
                metadata={"path": str(filepath), "format": "docx"},
            )
        except Exception as exc:
            return ToolResult(
                tool_name=self.name,
                status=ToolStatus.ERROR,
                error=f"DOCX generation failed: {exc}",
            )

    async def _generate_xlsx(self, filename: str, content: dict) -> ToolResult:
        try:
            from openpyxl import Workbook
        except ImportError:
            return ToolResult(
                tool_name=self.name,
                status=ToolStatus.ERROR,
                error="openpyxl not installed. Install with: pip install openpyxl",
            )

        try:
            wb = Workbook()
            # Remove default sheet
            wb.remove(wb.active)

            sheets = content.get("sheets", [])
            if not sheets:
                sheets = [{"name": "Sheet1", "headers": [], "rows": []}]

            for sheet_data in sheets:
                ws = wb.create_sheet(title=sheet_data.get("name", "Sheet"))
                headers = sheet_data.get("headers", [])
                if headers:
                    ws.append(headers)
                for row in sheet_data.get("rows", []):
                    ws.append(row)

            output_dir = self._ensure_output_dir()
            filepath = output_dir / f"{filename}.xlsx"
            wb.save(str(filepath))

            return ToolResult(
                tool_name=self.name,
                status=ToolStatus.SUCCESS,
                output=f"Spreadsheet saved to {filepath}",
                metadata={"path": str(filepath), "format": "xlsx"},
            )
        except Exception as exc:
            return ToolResult(
                tool_name=self.name,
                status=ToolStatus.ERROR,
                error=f"XLSX generation failed: {exc}",
            )

    async def _generate_pdf(self, filename: str, content: dict) -> ToolResult:
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        except ImportError:
            return ToolResult(
                tool_name=self.name,
                status=ToolStatus.ERROR,
                error="reportlab not installed. Install with: pip install reportlab",
            )

        try:
            output_dir = self._ensure_output_dir()
            filepath = output_dir / f"{filename}.pdf"

            doc = SimpleDocTemplate(str(filepath), pagesize=letter)
            styles = getSampleStyleSheet()
            story = []

            title = content.get("title")
            if title:
                story.append(Paragraph(title, styles["Title"]))
                story.append(Spacer(1, 12))

            for para in content.get("paragraphs", []):
                story.append(Paragraph(para, styles["BodyText"]))
                story.append(Spacer(1, 6))

            doc.build(story)

            return ToolResult(
                tool_name=self.name,
                status=ToolStatus.SUCCESS,
                output=f"PDF saved to {filepath}",
                metadata={"path": str(filepath), "format": "pdf"},
            )
        except Exception as exc:
            return ToolResult(
                tool_name=self.name,
                status=ToolStatus.ERROR,
                error=f"PDF generation failed: {exc}",
            )
